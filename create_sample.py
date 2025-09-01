"""Create sample DOCX file."""
from docx import Document

doc = Document()
doc.add_heading('샘플 계약서', 0)

doc.add_heading('제1조 (목적)', level=1)
doc.add_paragraph('본 계약은 소프트웨어 개발 용역에 관한 사항을 정함을 목적으로 한다.')

doc.add_heading('제2조 (대금 지급)', level=1)
doc.add_paragraph('발주자는 계약 대금을 다음과 같이 지급한다.')
doc.add_paragraph('지연 시 이자를 지급한다.')

doc.add_heading('제3조 (검수)', level=1)
doc.add_paragraph('검수요청일로부터 5영업일 내 결과 통지, 미통지는 합격으로 본다.')
doc.add_paragraph('재검수는 3영업일 내 수행한다.')

doc.add_heading('제4조 (인력 관리)', level=1)
doc.add_paragraph('프로젝트 참여 인력의 고용에 관한 사항은 별도 협의한다.')

doc.add_heading('제5조 (손해배상)', level=1)
doc.add_paragraph('손해배상 책임은 무제한으로 한다.')

doc.save('sample_docs/sample1.docx')
print("Sample DOCX created successfully!")
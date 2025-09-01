"""Create realistic sample DOCX file."""
from docx import Document

doc = Document()
doc.add_heading('IT 서비스 용역 계약서', 0)

doc.add_paragraph('㈜가나다(이하 "위탁인"이라 한다)와 ㈜갑을병(이하 "수탁인"이라 한다)은 다음과 같이 IT 서비스 용역 계약(이하 "본 계약"이라 한다)을 체결한다.')

doc.add_heading('제1조 (목적)', level=1)
doc.add_paragraph('본 계약은 소프트웨어 개발 용역에 관한 사항을 정함을 목적으로 한다.')

doc.add_heading('제2조 (계약기간)', level=1)
doc.add_paragraph('본 계약의 기간은 2024년 1월 1일부터 2024년 12월 31일까지 12개월로 한다.')

doc.add_heading('제3조 (계약금액)', level=1)
doc.add_paragraph('본 계약의 총 계약금액은 50,000,000원으로 한다.')

doc.add_heading('제4조 (대금 지급)', level=1)
doc.add_paragraph('위탁인은 계약 대금을 다음과 같이 지급한다.')
doc.add_paragraph('지연 시 이자를 지급한다.')

doc.add_heading('제5조 (검수)', level=1)
doc.add_paragraph('검수요청일로부터 5영업일 내 결과 통지, 미통지는 합격으로 본다.')
doc.add_paragraph('재검수는 3영업일 내 수행한다.')

doc.add_heading('제6조 (인력 관리)', level=1)
doc.add_paragraph('프로젝트 참여 인력의 고용에 관한 사항은 별도 협의한다.')

doc.add_heading('제7조 (손해배상)', level=1)
doc.add_paragraph('손해배상 책임은 무제한으로 한다.')

doc.save('sample_docs/realistic_sample.docx')
print("Realistic sample DOCX created successfully!")
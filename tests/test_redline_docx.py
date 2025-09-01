"""Tests for redline DOCX generation."""
import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent.parent))

from docx import Document
import io
from app.services.redline_docx import RedlineGenerator
from app.models.schema import AuditResult, RuleResult


def test_redline_generation():
    """Test redline DOCX generation with missing and insufficient rules."""
    # Create sample document
    doc = Document()
    doc.add_paragraph("샘플 계약서 내용입니다.")
    doc.add_paragraph("지연 시 이자를 지급한다.")
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    original_bytes = doc_bytes.getvalue()
    
    # Create audit result with violations
    rule_results = [
        RuleResult(
            rule_id="PAY-001",
            title="지연이자 연12%",
            severity="required",
            status="insufficient",
            evidence="...지연 시 이자를 지급한다...",
            suggestion="발주자 지연 시 연 12% 가산이자를 지급한다.",
            redline_action="replace"
        ),
        RuleResult(
            rule_id="DAM-001",
            title="손해배상 상한=계약금액",
            severity="recommended",
            status="missing",
            suggestion="손해배상액의 한도는 계약금액으로 한다.",
            redline_action="add"
        )
    ]
    
    audit_result = AuditResult(
        total_rules=2,
        required_violations=1,
        recommended_violations=1,
        rule_results=rule_results
    )
    
    # Generate redline
    generator = RedlineGenerator()
    redline_bytes = generator.generate_redline_docx(original_bytes, audit_result)
    
    # Verify redline document
    redline_doc = Document(io.BytesIO(redline_bytes))
    doc_text = "\n".join([p.text for p in redline_doc.paragraphs])
    
    assert "자동 제안 (레드라인 요약)" in doc_text
    assert "제안" in doc_text  # suggestion already contains prefix
    assert "지연이자" in doc_text or "PAY-001" in doc_text
    assert "손해배상" in doc_text or "DAM-001" in doc_text
    
    print("[OK] Redline generation test passed")
    return redline_bytes


if __name__ == "__main__":
    redline_bytes = test_redline_generation()
    
    # Save test output
    with open("tests/test_redline_output.docx", "wb") as f:
        f.write(redline_bytes)
    
    print("[SUCCESS] Redline DOCX test passed!")
    print("[INFO] Test output saved to: tests/test_redline_output.docx")
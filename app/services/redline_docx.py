"""Generate redlined DOCX with suggestions."""
import io
from typing import List
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_UNDERLINE
from app.models.schema import AuditResult, RuleResult


class RedlineGenerator:
    def __init__(self):
        self.red_color = RGBColor(255, 0, 0)
    
    def generate_redline_docx(self, original_bytes: bytes, audit_result: AuditResult) -> bytes:
        """Generate redlined DOCX with suggestions."""
        # Load original document
        doc = Document(io.BytesIO(original_bytes))
        
        # Try inline suggestions first
        if hasattr(__import__('config'), 'INLINE_SUGGESTION_TRY') and __import__('config').INLINE_SUGGESTION_TRY:
            self._add_inline_suggestions(doc, audit_result.rule_results)
        
        # Add summary section at the end
        self._add_summary_section(doc, audit_result)
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    
    def _add_inline_suggestions(self, doc: Document, rule_results: List[RuleResult]):
        """Try to add inline suggestions near evidence."""
        for result in rule_results:
            if result.status == "present" or not result.evidence:
                continue
            
            # Simple approach: find paragraph containing evidence
            evidence_clean = result.evidence.replace("...", "").strip()[:20]
            
            for paragraph in doc.paragraphs:
                if evidence_clean.lower() in paragraph.text.lower():
                    # Add suggestion at end of paragraph (suggestion already contains prefix)
                    run = paragraph.add_run(f" [{result.suggestion}]")
                    run.font.color.rgb = self.red_color
                    run.underline = WD_UNDERLINE.SINGLE
                    break
    
    def _add_summary_section(self, doc: Document, audit_result: AuditResult):
        """Add improved redline summary section at document end."""
        # Add section break
        doc.add_page_break()
        
        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("자동 제안 (레드라인 요약)")
        title_run.bold = True
        title_run.font.size = 14
        
        # Summary stats
        doc.add_paragraph(f"총 {audit_result.total_rules}개 규칙 검토")
        doc.add_paragraph(f"필수 위반: {audit_result.required_violations}건")
        doc.add_paragraph(f"권고 위반: {audit_result.recommended_violations}건")
        doc.add_paragraph()
        
        # Group and sort suggestions
        violations = [r for r in audit_result.rule_results if r.status != "present"]
        
        # Sort by severity (required first) then by category
        def sort_key(result):
            severity_order = {"required": 0, "recommended": 1}
            category_order = {"PAY": 0, "ACC": 1, "HR": 2, "DEL": 3, "SVC": 4, "TAX": 5, "CHG": 6, "IP": 7}
            category = result.rule_id.split('-')[0]
            return (severity_order[result.severity], category_order.get(category, 99))
        
        sorted_violations = sorted(violations, key=sort_key)
        
        # Group by category
        current_category = None
        category_names = {
            "PAY": "대금 지급 관련",
            "ACC": "검수 및 인수 관련", 
            "HR": "인력 관리 관련",
            "DEL": "지연 및 납기 관련",
            "SVC": "용역 계약 관련",
            "TAX": "세무 및 원천징수 관련",
            "CHG": "변경관리 관련",
            "IP": "지식재산권 관련"
        }
        
        for result in sorted_violations:
            category = result.rule_id.split('-')[0]
            
            # Add category header if changed
            if current_category != category:
                current_category = category
                category_para = doc.add_paragraph()
                category_run = category_para.add_run(f"\n● {category_names.get(category, category)} 조항")
                category_run.bold = True
                category_run.font.size = 12
            
            # Add suggestion
            para = doc.add_paragraph()
            
            # Rule title and suggestion
            content_run = para.add_run(f"  • {result.title}: {result.suggestion}")
            content_run.font.color.rgb = self.red_color
            content_run.underline = WD_UNDERLINE.SINGLE
            
            # Evidence if available (smaller text)
            if result.evidence:
                evidence_para = doc.add_paragraph()
                evidence_run = evidence_para.add_run(f"    근거: {result.evidence}")
                evidence_run.font.size = 9
                evidence_run.italic = True